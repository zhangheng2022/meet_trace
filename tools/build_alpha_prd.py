from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "研会_AI_Alpha_PRD_重设计版.docx"

# Resolved preset: standard_business_brief
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B4D"
MUTED = "667085"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D0D5DD"
WHITE = "FFFFFF"
GREEN = "237A57"
AMBER = "8A6116"
RED = "9B1C1C"

LATIN_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    latin: str = LATIN_FONT,
    cjk: str = CJK_FONT,
):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cjk)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size: float, color: str, bold: bool = False):
    style.font.name = LATIN_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)


def set_cell_margins(cell, *, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, *, color=MID_GRAY, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)
        tag.set(qn("w:space"), "0")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA):
    assert sum(widths_dxa) == CONTENT_WIDTH_DXA
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, **CELL_MARGINS_DXA)
            set_cell_border(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_paragraph_shading(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_left_border(paragraph, color: str, size="18", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)


def next_ids(numbering_part):
    root = numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in root.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in root.findall(qn("w:num"))]
    return (max(abstract_ids, default=0) + 1, max(num_ids, default=0) + 1)


def add_numbering_definition(doc: Document, *, fmt: str, text: str) -> int:
    numbering_part = doc.part.numbering_part
    abstract_id, num_id = next_ids(numbering_part)
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), CJK_FONT)
    r_pr.append(r_fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering_part.element.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering_part.element.append(num)
    return num_id


def apply_num(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_text(paragraph, text: str, *, bold=None, color=None, size=None, italic=None):
    run = paragraph.add_run(text)
    if any(value is not None for value in (bold, color, size, italic)):
        set_run_font(run, bold=bold, color=color, size=size, italic=italic)
    return run


def add_para(doc, text="", *, style=None, before=0, after=None, keep=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_together = keep
    add_text(p, text)
    return p


def add_bullets(doc, items: list[str], bullet_num_id: int):
    for item in items:
        p = doc.add_paragraph(style="Alpha Bullet")
        apply_num(p, bullet_num_id)
        add_text(p, item)


def add_numbered(doc, items: list[str], decimal_num_id: int):
    for item in items:
        p = doc.add_paragraph(style="Alpha Numbered")
        apply_num(p, decimal_num_id)
        add_text(p, item)


def add_callout(doc, label: str, body: str, *, fill=LIGHT_BLUE, accent=BLUE):
    p = doc.add_paragraph(style="Decision Callout")
    set_paragraph_shading(p, fill)
    set_paragraph_left_border(p, accent)
    add_text(p, f"{label}  ", bold=True, color=accent)
    add_text(p, body, color=INK)
    return p


def add_kv_table(doc, rows: list[tuple[str, str]], *, label_width=2100):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [label_width, CONTENT_WIDTH_DXA - label_width])
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(("项目", "说明")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.style = "Table Header"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        add_text(p, header)
    for label, value in rows:
        row = table.add_row()
        for cell in row.cells:
            set_cell_margins(cell, **CELL_MARGINS_DXA)
            set_cell_border(cell)
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        p0 = row.cells[0].paragraphs[0]
        p0.style = "Table Header"
        p0.paragraph_format.space_after = Pt(0)
        add_text(p0, label)
        p1 = row.cells[1].paragraphs[0]
        p1.style = "Table Body"
        p1.paragraph_format.space_after = Pt(0)
        add_text(p1, value)
    set_table_geometry(table, [label_width, CONTENT_WIDTH_DXA - label_width])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_matrix(doc, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.style = "Table Header"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(header) <= 6 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        add_text(p, header)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.style = "Table Body"
            p.paragraph_format.space_after = Pt(0)
            if idx == 0 and len(value) <= 10:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text(p, value)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_requirement(doc, req_id: str, title: str, behavior: list[str], acceptance: list[str]):
    p = doc.add_paragraph(style="Heading 3")
    add_text(p, f"{req_id}  {title}")
    add_text(doc.add_paragraph("产品行为", style="Requirement Label"), "：")
    add_bullets(doc, behavior, BULLET_NUM_ID)
    add_text(doc.add_paragraph("通过条件", style="Requirement Label"), "：")
    add_bullets(doc, acceptance, BULLET_NUM_ID)


def configure_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal, size=11, color="000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    set_style_font(title, size=25, color=INK, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, size=13, color=MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle.paragraph_format.keep_with_next = True

    h1 = styles["Heading 1"]
    set_style_font(h1, size=16, color=BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    set_style_font(h2, size=13, color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    set_style_font(h3, size=12, color=DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    for name in [
        "Kicker",
        "Meta",
        "Decision Callout",
        "Alpha Bullet",
        "Alpha Numbered",
        "Requirement Label",
        "Table Header",
        "Table Body",
        "Small Muted",
    ]:
        if name not in styles:
            styles.add_style(name, 1)

    kicker = styles["Kicker"]
    set_style_font(kicker, size=10, color=BLUE, bold=True)
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(6)
    kicker.paragraph_format.keep_with_next = True

    meta = styles["Meta"]
    set_style_font(meta, size=9.5, color=MUTED)
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(2)

    callout = styles["Decision Callout"]
    set_style_font(callout, size=10.5, color=INK)
    callout.paragraph_format.left_indent = Inches(0.14)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(8)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.line_spacing = 1.15
    callout.paragraph_format.keep_together = True

    for name in ("Alpha Bullet", "Alpha Numbered"):
        style = styles[name]
        set_style_font(style, size=11, color="000000")
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    req = styles["Requirement Label"]
    set_style_font(req, size=9.5, color=MUTED, bold=True)
    req.paragraph_format.space_before = Pt(2)
    req.paragraph_format.space_after = Pt(2)
    req.paragraph_format.keep_with_next = True

    table_header = styles["Table Header"]
    set_style_font(table_header, size=9.2, color=INK, bold=True)
    table_header.paragraph_format.space_before = Pt(0)
    table_header.paragraph_format.space_after = Pt(0)
    table_header.paragraph_format.line_spacing = 1.10

    table_body = styles["Table Body"]
    set_style_font(table_body, size=9.2, color=INK)
    table_body.paragraph_format.space_before = Pt(0)
    table_body.paragraph_format.space_after = Pt(0)
    table_body.paragraph_format.line_spacing = 1.10

    small = styles["Small Muted"]
    set_style_font(small, size=9, color=MUTED)
    small.paragraph_format.space_before = Pt(0)
    small.paragraph_format.space_after = Pt(4)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    add_text(p, "研会 AI  |  Alpha 产品需求", size=9, color=MUTED, bold=True)
    add_text(p, "\tV0.2 · 2026-07-23", size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    add_text(fp, "内部评审  ·  第 ", size=9, color=MUTED)
    add_page_field(fp)
    add_text(fp, " 页", size=9, color=MUTED)


def page_break(doc):
    doc.add_page_break()


doc = Document()
configure_styles(doc)
for section in doc.sections:
    configure_section(section)

BULLET_NUM_ID = add_numbering_definition(doc, fmt="bullet", text="•")
DECIMAL_NUM_ID = add_numbering_definition(doc, fmt="decimal", text="%1.")

# Opening block: memo_masthead
p = doc.add_paragraph(style="Kicker")
add_text(p, "产品需求文档  /  两周封闭 Alpha")
p = doc.add_paragraph(style="Title")
add_text(p, "研会 AI")
p = doc.add_paragraph(style="Subtitle")
add_text(p, "安卓线下会议实时转录与 AI 总结")

meta_rows = [
    ("版本", "V0.2（重设计版）"),
    ("状态", "范围已确认，可进入技术可行性检查与两周开发"),
    ("目标平台", "Android；首轮仅覆盖选定测试机型"),
    ("验证周期", "2 周；2 名工程师 + 1 名兼职产品/设计"),
    ("试用规模", "5–10 名封闭用户，至少 20 场真实线下会议"),
]
for label, value in meta_rows:
    p = doc.add_paragraph(style="Meta")
    add_text(p, f"{label}：", bold=True, color=INK, size=9.5)
    add_text(p, value)

add_callout(
    doc,
    "产品决策",
    "本地录音是唯一事实源；实时转录是允许降级的临时结果；会议结束后基于完整音频生成最终转录，再生成统一结构的 AI 总结。",
)

doc.add_heading("一页结论", level=1)
add_kv_table(
    doc,
    [
        ("服务对象", "需要记录线下面对面会议的个人用户；会议主题不限。"),
        ("核心闭环", "开始会议 → 本地录音与实时转录 → 结束会议 → 最终转录与说话人分离 → AI 总结 → 核验、编辑、分享。"),
        ("P0 价值", "即使断网或实时转录失败，也不丢录音；会后仍能获得可核验、可编辑的转录与总结。"),
        ("明确不做", "iOS、线上会议接入、团队空间、项目管理、独立待办、声纹识别、离线 AI、私有化与外部系统集成。"),
        ("Alpha 判定", "20 场真实会议中至少 90% 完成全链路；至少 70% 的总结只需轻微修改即可使用。"),
    ],
)

doc.add_heading("1. 产品定义", level=1)
doc.add_heading("1.1 用户问题", level=2)
add_para(
    doc,
    "线下面对面会议中，参会者需要同时倾听、表达和记录。手工笔记常遗漏上下文，录音又难以回看；会议结束后，整理逐字稿和结论通常需要二次投入。",
)
add_para(
    doc,
    "研会 AI 用安卓手机完成稳定录音，在会中提供实时文字，在会后生成经过完整音频校正的最终转录与结构化总结，让用户能快速回顾、核验和分享。",
)

doc.add_heading("1.2 核心假设", level=2)
add_bullets(
    doc,
    [
        "用户愿意在普通线下会议中把安卓手机放在桌面中央并持续录音。",
        "实时文字能够降低遗漏感，但用户接受它是临时结果，且在弱网时允许暂停。",
        "用户更在意会后总结是否可靠、可核验，而不是摘要文风是否华丽。",
        "说话人分离、证据跳转和可编辑能力能够提升用户对 AI 总结的信任。",
        "使用成熟云端转录与大模型服务，能够在两周内完成可测试闭环。",
    ],
    BULLET_NUM_ID,
)

doc.add_heading("1.3 产品原则", level=2)
add_matrix(
    doc,
    ["原则", "产品含义", "实现约束"],
    [
        ["录音优先", "任何网络或转录故障都不能中断本地录音。", "音频增量落盘；实时链路与录音链路解耦。"],
        ["临时与最终分离", "会中字幕可变化，会后文本才是最终版本。", "界面明确标记状态；总结只使用最终转录。"],
        ["结论可核验", "关键总结必须能回到原文与录音。", "每条关键项绑定转录片段和时间戳。"],
        ["AI 不补事实", "缺失的负责人、时间或结论必须留空。", "提示词、结构校验与编辑流程共同约束。"],
        ["Alpha 先验证", "先证明用户价值与链路可靠性，再扩展协作。", "不为团队空间、项目管理和生态集成预留复杂交互。"],
    ],
    [1500, 3900, 3960],
)

doc.add_heading("2. 目标与边界", level=1)
doc.add_heading("2.1 两周目标", level=2)
add_numbered(
    doc,
    [
        "完成一条在真实安卓设备上可连续使用的端到端会议记录链路。",
        "验证实时转录与会后最终转录能够共存，并在弱网条件下正确降级与恢复。",
        "验证统一 AI 总结对通用会议是否足够有用、可信且容易修正。",
        "收集真实用户对录音意愿、总结使用率、修改量和失败原因的数据。",
    ],
    DECIMAL_NUM_ID,
)

doc.add_heading("2.2 非目标", level=2)
add_bullets(
    doc,
    [
        "不追求公开发布、应用商店上架或大规模设备兼容。",
        "不支持 iOS、平板、桌面端或 Web 端录音。",
        "不接入腾讯会议、飞书、钉钉、Zoom 等线上会议系统，也不录制系统内部音频。",
        "不建设团队空间、成员权限、项目管理、协同编辑或独立待办系统。",
        "不做声纹注册和真实姓名自动识别；只提供说话人编号与会后重命名。",
        "不做离线转录、端侧大模型、私有化部署和第三方业务系统集成。",
        "不承诺超过 30 分钟会议的稳定性验收；更长会议可试用但不作为 Alpha 通过条件。",
    ],
    BULLET_NUM_ID,
)

doc.add_heading("2.3 发布范围", level=2)
add_matrix(
    doc,
    ["层级", "进入 Alpha", "后移"],
    [
        ["平台", "Android；选定测试机型；封闭安装包", "iOS、应用商店发布、多端同步"],
        ["采集", "麦克风录音、增量本地保存、30 分钟稳定性", "系统音频、外接录音设备深度适配"],
        ["转录", "实时临时文本、会后最终文本、说话人分离与重命名", "声纹身份、实时翻译、专业词库管理"],
        ["AI", "统一五段式总结、证据绑定、编辑与确认", "多会议模板、问答、跨会议知识库"],
        ["输出", "系统分享纯文本/Markdown", "公开分享链接、PDF、协作评论、外部集成"],
        ["账号", "封闭测试账号或邀请登录", "公开注册、找回流程、组织和权限"],
    ],
    [1300, 4030, 4030],
)

page_break(doc)
doc.add_heading("3. 核心用户闭环", level=1)
add_numbered(
    doc,
    [
        "进入会议列表，点击“开始新会议”。",
        "看到录音与云端处理说明，确认已取得参会者同意；授予麦克风权限。",
        "输入可选会议标题并开始。App 立即在本地增量保存音频，同时上传音频片段供实时转录。",
        "会中查看带“临时”标记的实时文字。弱网时文字可暂停，但录音状态必须保持正常。",
        "点击结束后，App 封存本地音频并补传缺失片段；服务端生成最终转录和说话人分离结果。",
        "系统只基于最终转录生成统一结构的 AI 总结。",
        "用户重命名说话人，点击关键总结跳回原文并播放对应录音，随后编辑或删除错误内容。",
        "用户确认结果，通过安卓系统分享面板导出纯文本或 Markdown；也可删除整场会议。",
    ],
    DECIMAL_NUM_ID,
)

add_callout(
    doc,
    "失败边界",
    "实时转录失败不等于会议失败。只有本地录音无法持续写入、最终音频不可恢复，才判定本场会议记录失败。",
    fill="FFF6E5",
    accent=AMBER,
)

doc.add_heading("4. 状态模型与关键规则", level=1)
add_matrix(
    doc,
    ["对象", "状态", "用户看到什么", "系统规则"],
    [
        ["会议", "录音中", "计时、录音正常标识、实时文字区域", "持续本地落盘；音频片段异步上传。"],
        ["实时转录", "正常", "临时文本持续出现", "仅用于会中阅读，不触发最终总结。"],
        ["实时转录", "已降级", "“网络不佳，仍在录音”", "停止或延迟上传；网络恢复后补传。"],
        ["会议", "处理中", "上传、最终转录、总结的分阶段进度", "任一阶段失败可重试，不能丢失本地音频。"],
        ["最终转录", "可用", "稳定文本、发言人编号、时间戳", "覆盖临时文本；后续总结以此版本为准。"],
        ["AI 总结", "待确认", "五段式结果与证据入口", "允许编辑、删除；不自动标记为用户确认。"],
        ["会议", "已完成", "确认后的转录、总结、分享入口", "保留云端文本；音频按保留策略清理。"],
        ["会议", "处理失败", "失败原因、重试和保留录音说明", "允许重新上传、重新转录或重新总结。"],
    ],
    [1350, 1450, 3000, 3560],
)

doc.add_heading("5. P0 功能需求", level=1)
doc.add_heading("5.1 账号与会议入口", level=2)
add_requirement(
    doc,
    "FR-001",
    "封闭测试登录",
    [
        "仅支持测试账号或邀请方式进入；不提供公开注册、组织加入和复杂账号管理。",
        "登录后进入个人会议列表，按时间倒序展示标题、日期、时长和处理状态。",
    ],
    [
        "5–10 名测试用户可以独立进入自己的会议列表。",
        "用户不能看到其他测试用户的会议。",
    ],
)
add_requirement(
    doc,
    "FR-002",
    "开始会议",
    [
        "用户可输入可选标题；为空时按日期时间生成标题。",
        "首次录音前展示数据处理和参会者同意提示，并请求麦克风权限。",
        "权限被拒绝时明确说明原因并提供再次授权入口。",
    ],
    [
        "从会议列表到开始录音不超过 3 个主要操作。",
        "未经麦克风授权不能进入录音状态。",
    ],
)

doc.add_heading("5.2 本地录音与实时转录", level=2)
add_requirement(
    doc,
    "FR-010",
    "本地增量录音",
    [
        "开始后立即采集麦克风音频，并以可恢复方式持续写入 App 私有目录。",
        "录音链路不得依赖网络、实时转录或 AI 服务。",
        "录音页持续显示计时、录音状态、剩余空间预警和结束入口。",
        "Alpha 允许屏幕常亮或前台运行；后台与锁屏录音若未通过测试，不得宣称支持。",
    ],
    [
        "选定测试机型上连续录制 30 分钟不崩溃、不丢失可播放音频。",
        "App 被异常终止后，重新打开能发现并恢复未完成会议。",
        "存储空间不足时提前提示，并在无法继续时安全封存已有音频。",
    ],
)
add_requirement(
    doc,
    "FR-011",
    "实时临时转录",
    [
        "录音期间按片上传音频并展示实时文字；界面明确标记为“临时转录”。",
        "正常网络下，实时首字延迟 P95 不超过 3 秒。",
        "临时文本允许被后续识别结果修订，不作为分享与 AI 总结的最终来源。",
    ],
    [
        "实时文字按时间顺序追加，用户可在会中滚动查看。",
        "网络正常时，20 场测试会议的实时延迟达到目标。",
        "实时服务失败时，录音仍持续且界面不误报整场失败。",
    ],
)
add_requirement(
    doc,
    "FR-012",
    "弱网降级与补传",
    [
        "网络断开或服务不可用时，页面显示“实时转录暂停，录音仍在继续”。",
        "网络恢复后自动上传本地缺失片段，不重复创建转录内容。",
        "用户结束会议时，未上传片段继续进入会后补传队列。",
    ],
    [
        "模拟断网 3 分钟后恢复，录音完整且缺失片段可补传。",
        "补传后最终转录覆盖断网时段。",
        "用户始终能区分录音状态和转录状态。",
    ],
)

doc.add_heading("5.3 会后最终转录", level=2)
add_requirement(
    doc,
    "FR-020",
    "结束与处理",
    [
        "用户结束会议后，系统先封存本地音频，再展示上传、最终转录、AI 总结三个处理阶段。",
        "最终转录必须基于完整音频重新生成，不直接把实时临时文本视为最终稿。",
        "各阶段可独立重试；重新总结不得强制重新转录。",
    ],
    [
        "30 分钟会议在结束后 5 分钟内产出最终转录和 AI 总结。",
        "处理失败时显示具体失败阶段和可执行的重试入口。",
    ],
)
add_requirement(
    doc,
    "FR-021",
    "说话人分离与重命名",
    [
        "最终转录按发言段显示“发言人 1 / 2 / 3”等编号。",
        "用户可将同一编号批量重命名，修改同步到全文和总结展示。",
        "系统不注册声纹、不推断真实身份，也不跨会议复用身份。",
    ],
    [
        "至少两人会议能产生可区分的发言段。",
        "一次重命名可更新该说话人在整场会议中的所有标签。",
    ],
)

doc.add_heading("5.4 AI 总结", level=2)
add_requirement(
    doc,
    "FR-030",
    "统一五段式总结",
    [
        "所有会议使用同一输出结构：会议概述、核心讨论、已确认结论、行动项、未决问题。",
        "行动项字段为事项、负责人、截止时间；原文未明确时负责人或截止时间保持为空。",
        "总结只使用最终转录，不引用会中临时文本。",
    ],
    [
        "五个区块始终存在；无内容时显示“未识别到”，不得编造填充。",
        "同一次输入重复生成时，字段结构保持稳定。",
    ],
)
add_requirement(
    doc,
    "FR-031",
    "证据绑定与事实约束",
    [
        "核心讨论、结论、行动项和未决问题中的每一条内容至少绑定一个转录片段与时间戳。",
        "模型不得补全未出现的人名、日期、数字和承诺。",
        "当不同发言相互矛盾且未形成结论时，应放入“未决问题”，而不是强行生成结论。",
    ],
    [
        "点击关键项可定位到对应转录片段。",
        "抽检中不存在无法追溯到原文的关键结论。",
        "负责人和截止时间不明确时保持空值。",
    ],
)
add_requirement(
    doc,
    "FR-032",
    "编辑、确认与重新生成",
    [
        "用户可修改或删除任意总结条目，也可重新生成整份总结。",
        "系统区分“AI 生成”和“用户已确认”状态；分享前不强制确认，但应清楚提示。",
        "重新生成前需提醒可能覆盖未保存的编辑。",
    ],
    [
        "编辑后退出再进入，修改内容仍然存在。",
        "删除条目不会删除原始转录。",
        "用户能明确判断当前结果是否已确认。",
    ],
)

doc.add_heading("5.5 核验、分享与删除", level=2)
add_requirement(
    doc,
    "FR-040",
    "原文与录音核验",
    [
        "转录段显示时间戳与说话人标签。",
        "点击总结证据跳到对应转录；点击转录片段可从对应时间播放录音。",
        "音频播放与文字高亮保持基本同步。",
    ],
    [
        "从任意关键总结项到对应录音不超过 2 个操作。",
        "定位误差不超过 3 秒。",
    ],
)
add_requirement(
    doc,
    "FR-041",
    "系统分享",
    [
        "通过安卓系统分享面板导出纯文本或 Markdown。",
        "默认包含标题、时间、五段式总结；用户可选择附加完整转录。",
        "Alpha 不生成公开网页链接，不处理多人访问权限。",
    ],
    [
        "结果可成功分享到至少一种测试应用。",
        "分享文本不包含内部 ID、模型提示词或调试信息。",
    ],
)
add_requirement(
    doc,
    "FR-042",
    "删除会议",
    [
        "用户可删除整场会议；操作前明确提示将删除录音、转录和总结。",
        "删除请求同步清理云端数据和本地缓存；失败时展示待重试状态。",
    ],
    [
        "删除后会议不再出现在列表中。",
        "云端记录和本地音频均无法继续访问。",
    ],
)

page_break(doc)
doc.add_heading("6. 页面与关键交互", level=1)
add_matrix(
    doc,
    ["页面", "主要内容", "必须处理的状态"],
    [
        ["会议列表", "新建入口、会议标题、时间、时长、处理状态", "空列表、处理中、失败、已完成"],
        ["开始会议", "可选标题、录音同意说明、麦克风权限", "首次授权、拒绝、存储不足"],
        ["录音页", "计时、录音状态、临时转录、网络状态、结束", "正常、弱网降级、转录失败、录音故障"],
        ["处理页", "上传、最终转录、AI 总结三段进度", "进行中、部分失败、重试、完成"],
        ["结果页", "五段式总结、最终转录、说话人重命名、证据与播放", "待确认、已编辑、已确认、重新生成"],
        ["分享/删除", "系统分享选项、是否附原文、删除确认", "分享失败、删除待重试"],
    ],
    [1550, 4110, 3700],
)

doc.add_heading("6.1 录音页信息优先级", level=2)
add_bullets(
    doc,
    [
        "第一优先级：录音是否仍在进行。使用持续可见的红色录音标识与计时。",
        "第二优先级：本地文件是否安全。异常时给出明确、不可忽略的错误。",
        "第三优先级：实时转录是否正常。弱网提示不能覆盖或混淆录音状态。",
        "第四优先级：临时文字内容。允许延迟、修订和滚动查看。",
    ],
    BULLET_NUM_ID,
)

doc.add_heading("6.2 处理页进度规则", level=2)
add_bullets(
    doc,
    [
        "进度按“补传音频 → 最终转录 → AI 总结”展示，不使用无法解释的单一百分比。",
        "用户可离开处理页；会议列表持续显示状态。",
        "最终转录完成但总结失败时，用户仍可阅读和分享转录，并单独重试总结。",
        "本地音频未完成上传前不得自动清理。",
    ],
    BULLET_NUM_ID,
)

doc.add_heading("7. AI 输出契约", level=1)
add_matrix(
    doc,
    ["区块", "输出要求", "证据规则"],
    [
        ["会议概述", "1–3 句话说明主题、目的和总体结果；无法判断目的时只描述讨论主题。", "至少绑定一个代表性片段。"],
        ["核心讨论", "按主题列出主要观点与分歧，不把观点写成已确认结论。", "每条至少一个片段；分歧可绑定多个片段。"],
        ["已确认结论", "只记录明确达成或由主持人确认的结论。", "每条必须有形成结论的原文片段。"],
        ["行动项", "事项、负责人、截止时间；未知字段为空。", "每条必须绑定提出或确认任务的片段。"],
        ["未决问题", "记录未回答问题、待确认信息和未形成一致的分歧。", "每条绑定问题或分歧片段。"],
    ],
    [1500, 5010, 2850],
)

doc.add_heading("7.1 AI 禁止行为", level=2)
add_bullets(
    doc,
    [
        "不得把建议、猜测、反问或个人观点自动升级为团队结论。",
        "不得补写原文不存在的负责人、日期、数字、原因或下一步。",
        "不得以常识替代会议内容，也不得把模型知识混入会议事实。",
        "不得生成没有证据时间戳的关键条目。",
        "不得在转录仍为临时状态时生成最终总结。",
    ],
    BULLET_NUM_ID,
)

doc.add_heading("7.2 质量评审方法", level=2)
add_bullets(
    doc,
    [
        "每场会议由使用者标记总结为“无需修改 / 轻微修改 / 大幅修改 / 不可用”。",
        "记录修改类型：事实错误、遗漏、错误结论、行动项字段错误、重复、表达问题。",
        "Alpha 以“无需修改 + 轻微修改”达到 70% 为可用门槛。",
        "证据可追溯率目标为 100%；任何无法回到原文的关键结论都视为严重缺陷。",
    ],
    BULLET_NUM_ID,
)

doc.add_heading("8. 数据、隐私与保留", level=1)
add_kv_table(
    doc,
    [
        ("处理方式", "音频上传云端完成实时/最终转录和 AI 总结；MVP 不提供离线或端侧处理。"),
        ("本地音频", "作为故障恢复副本保留；确认上传完成后默认保留 7 天，再自动清理。"),
        ("云端音频", "默认保留 7 天后自动删除；用户可提前删除整场会议；到期后录音播放不可用。"),
        ("文本数据", "转录和总结持续保留，直到用户删除会议。"),
        ("传输与存储", "网络传输使用 TLS；云端数据使用服务商提供的静态加密能力。"),
        ("录音同意", "开始前提示用户取得参会者同意；Alpha 不替代用户履行当地法律义务。"),
        ("最小化日志", "日志不得记录完整音频、完整转录、模型提示词中的敏感原文或访问令牌。"),
    ],
)

doc.add_heading("9. 非功能要求", level=1)
add_matrix(
    doc,
    ["类别", "Alpha 要求", "验收方式"],
    [
        ["可靠性", "本地录音与网络链路解耦；异常退出可恢复。", "选定机型进行 10 次 30 分钟录音与异常终止测试。"],
        ["性能", "实时首字延迟 P95 ≤ 3 秒；30 分钟会议会后 5 分钟内完成。", "记录端到端时间戳，不以主观感受代替数据。"],
        ["兼容性", "仅承诺 Android 10+ 的 3–5 台指定测试设备。", "形成设备清单并逐台执行主路径。"],
        ["安全", "TLS、访问隔离、删除能力、敏感日志最小化。", "账号隔离、网络抓包与删除回查。"],
        ["可观测性", "记录会议、上传、转录、总结各阶段成功/失败和耗时。", "每次失败能定位到阶段、错误码和重试结果。"],
        ["可访问性", "关键状态不能只依赖颜色；主要按钮有可读标签。", "人工检查录音、弱网和失败提示。"],
    ],
    [1450, 4690, 3220],
)

page_break(doc)
doc.add_heading("10. Alpha 指标与实验设计", level=1)
doc.add_heading("10.1 核心指标", level=2)
add_matrix(
    doc,
    ["指标", "目标", "计算口径"],
    [
        ["全链路完成率", "≥ 90%", "成功获得本地录音、最终转录和 AI 总结的会议数 / 已结束会议数。"],
        ["录音完整率", "100%", "抽检会议中可播放音频覆盖预期时长，且无不可恢复缺口。"],
        ["实时首字延迟 P95", "≤ 3 秒", "音频片段产生到首个临时文字出现的延迟。"],
        ["会后处理时长", "≤ 5 分钟", "30 分钟会议从结束到最终转录与总结均可用。"],
        ["总结可用率", "≥ 70%", "用户评价为“无需修改”或“轻微修改”的总结占比。"],
        ["关键项可追溯率", "100%", "带证据的结论、行动项和未决问题 / 全部关键项。"],
    ],
    [2200, 1500, 5660],
)

doc.add_heading("10.2 试用设计", level=2)
add_numbered(
    doc,
    [
        "招募 5–10 名安卓用户，覆盖至少 3 种手机型号和不同会议主题。",
        "每名用户完成 2–4 场线下会议，总计至少 20 场；目标时长 10–30 分钟。",
        "会议结束后立即采集总结可用度、修改类型和是否分享。",
        "每日查看失败漏斗：开始失败、录音失败、上传失败、最终转录失败、总结失败。",
        "第 10 天依据通过标准做继续、修复后继续或停止的决定。",
    ],
    DECIMAL_NUM_ID,
)

doc.add_heading("10.3 Go / No-Go", level=2)
add_callout(
    doc,
    "继续投入",
    "全链路完成率达到 90%，录音无不可恢复丢失，总结可用率达到 70%，且至少一半测试用户愿意继续记录下一场会议。",
    fill="EAF7F0",
    accent=GREEN,
)
add_callout(
    doc,
    "暂停扩展",
    "出现不可恢复录音丢失、最终转录长期不稳定，或总结可用率低于 50%。在这些问题解决前，不增加 iOS、团队协作或更多导出能力。",
    fill="FDECEC",
    accent=RED,
)

doc.add_heading("11. 两周交付计划", level=1)
add_matrix(
    doc,
    ["日程", "安卓工程", "后端 / AI", "产品与验证"],
    [
        ["Day 1", "录音技术尖峰、测试机型、权限流程", "实时/批量 ASR 与说话人分离服务验证", "锁定服务商、测试脚本与失败边界"],
        ["Day 2", "增量落盘、会议本地状态", "账号隔离、会议与音频上传接口", "开始页与录音页关键状态"],
        ["Day 3", "分片上传、实时文字 UI", "实时转录流与状态回传", "实时/录音状态文案验收"],
        ["Day 4", "断网队列、补传与恢复", "最终转录、说话人分离", "弱网与异常退出用例"],
        ["Day 5", "端到端串联", "五段式总结与证据映射", "内部完成 3 场冒烟测试"],
        ["Day 6", "结果页、说话人重命名", "证据定位与音频时间戳", "总结质量标注表"],
        ["Day 7", "编辑、系统分享、删除", "重试、数据保留与删除", "隐私提示与分享模板"],
        ["Day 8", "兼容性与恢复修复", "日志、指标、失败诊断", "3–5 台设备回归"],
        ["Day 9", "Pilot 包与高优缺陷修复", "容量与性能修复", "首轮真实用户试用"],
        ["Day 10", "发布候选包", "服务配置固化", "20 场数据复盘与 Go/No-Go"],
    ],
    [1050, 2810, 2810, 2690],
)

doc.add_heading("11.1 资源分工假设", level=2)
add_bullets(
    doc,
    [
        "工程师 A：Android 录音、状态持久化、上传队列和客户端页面。",
        "工程师 B：账号与会议接口、云端转录编排、AI 总结、证据映射和可观测性。",
        "兼职产品/设计：交互状态、验收用例、试用招募、数据标注和每日决策。",
        "必须使用成熟的云端实时 ASR、批量 ASR/说话人分离和大模型服务；两周内不自研模型。",
    ],
    BULLET_NUM_ID,
)

doc.add_heading("11.2 Day 1 强制退出条件", level=2)
add_bullets(
    doc,
    [
        "所选服务不能同时支持可接受的实时转录和会后最终转录。",
        "批量转录无法返回时间戳或说话人分离结果，导致证据链无法成立。",
        "目标网络环境下 30 分钟音频处理成本或速度不可接受。",
        "若触发退出条件，必须立即更换服务；若仍无法满足说话人分离与证据链要求，应发起范围复审，不得擅自取消已确认的 P0。",
    ],
    BULLET_NUM_ID,
)

doc.add_heading("12. 验收场景", level=1)
add_matrix(
    doc,
    ["编号", "场景", "预期结果"],
    [
        ["AT-01", "正常网络下完成 30 分钟会议", "实时文字出现；结束后 5 分钟内生成最终转录与总结；可核验、编辑和分享。"],
        ["AT-02", "录音中断网 3 分钟后恢复", "本地录音不中断；实时状态降级；恢复后补传；最终文本覆盖断网时段。"],
        ["AT-03", "实时转录服务返回错误", "页面提示转录暂停；录音继续；会后仍可完成最终转录。"],
        ["AT-04", "录音中强制终止 App", "重新打开后发现未完成会议，可恢复上传与后续处理，已有音频可播放。"],
        ["AT-05", "两人及以上交替发言", "最终转录显示发言人编号；批量重命名后全文同步更新。"],
        ["AT-06", "原文未提及负责人或截止时间", "行动项相应字段为空，不由 AI 补写。"],
        ["AT-07", "总结包含关键结论", "点击结论跳到正确转录并从相近时间播放录音，定位误差 ≤ 3 秒。"],
        ["AT-08", "用户删除会议", "列表消失；云端与本地数据均不可继续访问；失败时可重试。"],
    ],
    [1100, 3390, 4870],
)

doc.add_heading("13. 主要风险与应对", level=1)
add_matrix(
    doc,
    ["风险", "触发信号", "两周内应对"],
    [
        ["安卓录音不稳定", "音频缺段、文件损坏、进程终止", "限定测试机型；增量落盘；优先前台稳定性；每日 30 分钟压力测试。"],
        ["手机摆位导致音质差", "远端说话人漏字明显", "开始页给出桌面中央摆放提示；记录设备与环境；暂不承诺复杂远场。"],
        ["实时与最终文本混淆", "用户看到会后文字变化并认为丢失", "明确“临时/最终”标签；处理页解释重新转录；总结仅用最终稿。"],
        ["说话人分离不稳定", "频繁串人或单人被拆成多人", "允许批量重命名；把准确率作为观察项；不做身份推断。"],
        ["AI 生成错误结论", "总结与原文不符或补事实", "结构化输出、证据强制绑定、空值规则、用户确认和错误类型采集。"],
        ["云端隐私顾虑", "用户不愿录制或要求立即删除", "录音前说明、7 天音频保留、整场删除；Alpha 不承诺私有化。"],
        ["两周范围失控", "新增模板、团队、导出或双端需求", "所有新增项进入候选池；只有阻断核心闭环的缺陷可进入当前迭代。"],
    ],
    [2200, 2690, 4470],
)

doc.add_heading("14. 决策记录", level=1)
add_matrix(
    doc,
    ["决策", "结论", "原因"],
    [
        ["核心能力", "实时转录 + 会后最终转录 + AI 总结", "用户明确要求 A、B 同时存在。"],
        ["事实源", "本地录音", "保障弱网和服务故障下仍可恢复。"],
        ["平台", "Android only", "两周期限下避免双端音频工程风险。"],
        ["产品形态", "个人工具", "砍掉团队、项目和待办系统，聚焦核心闭环。"],
        ["会议范围", "通用主题、仅线下面对面会议", "保持用户场景广度，同时控制采集方式。"],
        ["总结模板", "统一五段式", "避免四类会议模板带来的产品与提示词膨胀。"],
        ["说话人", "自动编号 + 手动重命名", "提高可读性，不引入声纹身份系统。"],
        ["数据边界", "云端处理；音频 7 天；文本持续保留", "兼顾两周可交付性与用户控制。"],
        ["验证标准", "5–10 人、20 场、90% 全链路、70% 总结可用", "用真实使用结果而不是演示完成度判断。"],
    ],
    [2200, 2700, 4460],
)

doc.add_heading("15. Alpha 完成定义", level=1)
add_callout(
    doc,
    "Definition of Done",
    "一名受邀安卓用户可以在真实线下会议中稳定录音并查看实时临时文字；即使弱网也不丢录音；结束后 5 分钟内获得带说话人、时间戳和证据的最终转录及五段式 AI 总结；用户能够核验、编辑、分享或删除结果；20 场试用数据达到既定门槛。",
)

add_para(
    doc,
    "本版本替代原 PRD 中以完整产品形态为导向的 MVP 范围。任何 iOS、团队协作、项目管理、更多会议模板、外部集成或私有化需求，只有在 Alpha 通过 Go/No-Go 后才重新评估。",
    style="Small Muted",
    before=8,
    after=0,
)

# Document metadata
doc.core_properties.title = "研会 AI｜安卓线下会议实时转录与 AI 总结 Alpha PRD"
doc.core_properties.subject = "两周封闭 Alpha 产品需求"
doc.core_properties.author = "产品团队"
doc.core_properties.keywords = "会议转录, AI总结, Android, Alpha, PRD"
doc.core_properties.comments = "基于 grill-me 访谈重构"

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
